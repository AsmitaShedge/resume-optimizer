from docx import Document


def generate_improved_resume(
    name,
    email,
    phone,
    skills,
    missing_skills,
    original_text,
    output_file
):
    doc = Document()

    # Header
    doc.add_heading(name, level=1)

    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph(f"Phone: {phone}")

    # Professional Summary

    doc.add_heading(
        "Professional Summary",
        level=2
    )
    
    if skills:
        top_skills = ", ".join(
            skills[:5]
    )
    else:
        top_skills = (
            "Python, SQL, Data Analysis"
    )
    
    summary = (
        f"Professional with knowledge of "
        f"{top_skills}. "
        f"Skilled in problem-solving, "
        f"data-driven decision making, "
        f"and technical project development."
    )
    
    doc.add_paragraph(summary)

    # Skills
    doc.add_heading("Skills", level=2)

    if skills:
        for skill in skills:
            doc.add_paragraph(skill)
    else:
        doc.add_paragraph(
            "Python, SQL, Machine Learning, Data Analysis"
        )

    # Recommended Skills

    if missing_skills:
    
        doc.add_heading(
            "Recommended Skills For Target Role",
            level=2
        )
    
        for skill in missing_skills:
            doc.add_paragraph(
                f"• {skill}"
            )

   # Projects

    doc.add_heading(
        "Projects",
        level=2
    )
    
    project_keywords = [
        "student performance",
        "dashboard",
        "management system",
        "analysis",
        "application"
    ]
    
    project_found = False
    
    for line in original_text.split("\n"):

        line = line.strip()
    
        if len(line) < 5:
            continue
    
        if line.lower() == "projects":
            continue
    
        if any(
            keyword in line.lower()
            for keyword in project_keywords
        ):
    
            line = line.strip()
    
        if any(
            keyword in line.lower()
            for keyword in project_keywords
        ):
    
            doc.add_paragraph(
                f"• {line}"
            )
    
            project_found = True
    
    if not project_found:
    
        doc.add_paragraph(
            "• Developed technical projects demonstrating practical skills and problem-solving abilities."
        )

    doc.save(output_file)

    doc.add_heading("Experience", level=2)

    doc.add_paragraph(
        "Professional experience details from original resume."
    )
    
    doc.add_heading("Education", level=2)
    
    doc.add_paragraph(
        "Educational qualifications from original resume."
    )
    
    doc.add_heading("Certifications", level=2)
    
    doc.add_paragraph(
        "Relevant certifications and achievements."
    )